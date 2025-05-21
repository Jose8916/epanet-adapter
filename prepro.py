import tkinter as tk
from tkinter import messagebox, filedialog, ttk
import numpy as np
import matplotlib

import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import re
from docx import Document
from fpdf import FPDF
from pymoo.algorithms.moo.nsga2 import NSGA2
from pymoo.optimize import minimize
from pymoo.core.problem import Problem
import multiprocessing as mp
from epanettools import epanet2 as et
from queue import Empty
from scipy.optimize import differential_evolution
import pandas as pd
from datetime import datetime
from io import BytesIO
from docx.shared import Inches
import epynet as ep

matplotlib.use('TkAgg') 
class OptimizadorHidraulico:
    def __init__(self, network):
        self.network = network
        self.variables = {
            'diametros': [('pipe-1', 0.1, 2.0), ('pipe-2', 0.2, 1.5)],
            'tiempos_valvulas': [('valve-1', 1.0, 10.0)],
            'rampas_bomba': [('pump-1', 1.0, 60.0)]
        }
        
    def funcion_objetivo(self, x):
        diam = {var[0]: x[i] for i, var in enumerate(self.variables['diametros'])}
        t_valv = {var[0]: x[i] for i, var in enumerate(self.variables['tiempos_valvulas'])}
        t_rampa = {var[0]: x[i] for i, var in enumerate(self.variables['rampas_bomba'])}
        
        for pipe_id, value in diam.items():
            self.network.get_link(pipe_id).diameter = value
            
        resultados = self.simular_red(t_valv, t_rampa)
        
        return [
            resultados['costo_total'], 
            resultados['max_presion'], 
            resultados['energia']
        ]
        
    def optimizar_nsga2(self):
        problem = ProblemaHidraulico(self.network, self.variables)
        algorithm = NSGA2(pop_size=100)
        res = minimize(problem, algorithm, ('n_gen', 200))
        
        pd.DataFrame(res.X).to_csv('pareto_solutions.csv')
        self.graficar_frontes(res.F)
        
        return res
    
class WaterHammerAnalyzer:
    def __init__(self, inp_file):
        self.network = ep.Network(inp_file)
        self.critical_nodes = []
        self.pipes_data = {}

    def load_pipe_parameters(self):
        for pipe in self.network.pipes:
            self.pipes_data[pipe.uid] = {
                'L': pipe.length,
                'D': pipe.diameter,
                'v0': pipe.velocity,
                'start_node': pipe.from_node.uid,
                'end_node': pipe.to_node.uid
            }

    def identify_critical_nodes(self):
        slopes = []
        for pipe in self.network.pipes:
            start_node = pipe.from_node  
            end_node = pipe.to_node      
            slope = abs(start_node.elevation - end_node.elevation)/pipe.length
            slopes.append((pipe.uid, slope))
        
        self.critical_nodes = sorted(slopes, key=lambda x: x[1], reverse=True)[:3]
        return self.critical_nodes

    def joukowsky_pressure(self, delta_v):
        c = 1200  
        rho = 1000  
        return rho * c * delta_v

    def run_analysis(self, valve_close_time):
        results = {}
        for pipe_id, data in self.pipes_data.items():
            t, Hup, Hdown, Q = metodo_caracteristicas_avanzado(
                rho=1000, 
                c=1200,
                L=data['L'],
                D=data['D'],
                v_inicial=data['v0'],
                t_cierre_valvula=valve_close_time
            )
            
            rho = 1000 
            g = 9.81 
            presion_down = [h * rho * g for h in Hdown]
            
            results[pipe_id] = {
                'data': (t, presion_down),
                'params': data
            }
        return results

def modelo_inercial_rigido(
    rho, L, D, v_inicial, t_cierre_valvula, 
    t_total=5.0, dt=0.01, P_inicial=300000
):
    g = 9.81
    A = np.pi * (D**2) / 4
    Q0 = v_inicial * A
    f = 0.02  # Factor de fricción
    
    nsteps = int(t_total/dt) + 1
    t = np.linspace(0, t_total, nsteps)
    
    Q = np.zeros(nsteps)
    P = np.zeros(nsteps)
    
    Q[0] = Q0
    P[0] = P_inicial 
    
    for i in range(1, nsteps):
        if t[i] < t_cierre_valvula:
            factor = 1 - t[i]/t_cierre_valvula
        else:
            factor = 0.0
        
        dQdt = (factor * P_inicial - (f * L/(2*D) * rho * Q[i-1]**2)/A**2) / (rho * L/A)
        
        Q[i] = Q[i-1] + dQdt * dt
        P[i] = factor * P_inicial - (f * L/(2*D) * rho * Q[i]**2)/A**2
        
    return t, P/1000, Q

def metodo_caracteristicas_simple(rho, c, L, D, v_inicial=1.0, t_cierre_valvula=1.0, escenario="sin_bomba", flujo_bomba=None, presion_bomba=None, t_total=5.0, dt=0.01, H_up=30.0, H_down=25.0):

    A = np.pi*(D**2)/4.0

    Q0 = v_inicial * A

    if escenario == "con_bomba" and flujo_bomba and presion_bomba:
        Q_bomba = flujo_bomba / 3600.0
        Q0 = Q_bomba 

    nsteps = int(t_total/dt) + 1
    t_array = np.linspace(0, t_total, nsteps)

    Hup = np.zeros(nsteps)
    Hdown = np.zeros(nsteps)
    Qline = np.zeros(nsteps)

    Hup[0] = H_up
    Hdown[0] = H_down
    Qline[0] = Q0

    g = 9.81
    f = 0.02
    Beta = f * L / (2*g*D)
    alpha = c/(g*A)  

    for n in range(1, nsteps):
        t_now = t_array[n]

        Hup[n] = H_up

        if t_now < t_cierre_valvula:
            valve_factor = 1.0 - (t_now / t_cierre_valvula)
        else:
            valve_factor = 0.0  

        Q_old = Qline[n-1]
        Hd_old = Hdown[n-1]

        c_plus = Hup[n] - Beta*Q_old*abs(Q_old) + alpha*Q_old

        c_minus = Hd_old + Beta*Q_old*abs(Q_old) - alpha*Q_old

        Qc = (c_plus - c_minus)/(2*Beta + 1e-12)

        Q_new = valve_factor * Qc

        H_d = c_minus + alpha*Q_new - Beta*Q_new*abs(Q_new)

        Qline[n] = Q_new
        Hdown[n] = H_d

    return t_array, Hup, Hdown, Qline


class BombaCentrifuga:
    def __init__(self, node_id, H0, Q0, potencia, inercia=100.0):
        self.H0 = H0  
        self.Q0 = Q0 
        self.potencia = potencia  
        self.inercia = inercia  
        self.vel_angular = 0.0  
        self.estado = False 
        self.node_id = node_id  
        self.curva_qh = lambda Q: H0 - 0.2*H0*(Q/Q0)**2  
        self.estado = "detenida"
        self.vel_angular = 0.0
        self.tiempo_arranque = 5.0  
        self.tiempo_parada = 3.0
        
    def arrancar(self, dt):
        if not self.estado:
            torque = (self.potencia * 1000) / (self.vel_angular + 1e-5)
            self.vel_angular += (torque / self.inercia) * dt
            if self.vel_angular >= 1500 * (2*np.pi/60):  
                self.estado = True
                
    def parar(self, dt):
        if self.estado:
            torque_frenado = -self.vel_angular * 0.1 
            self.vel_angular += (torque_frenado / self.inercia) * dt
            if self.vel_angular <= 0:
                self.estado = False
                self.vel_angular = 0.0
                
    def curva_caracteristica(self, Q):
        return self.H0 - 0.2*self.H0*(Q/self.Q0)**2  

    def actualizar_estado(self, t, flujo_actual):
        if self.estado == "arrancando":
            self.vel_angular += (1500 - self.vel_angular)/self.tiempo_arranque
            if self.vel_angular >= 1450:
                self.estado = "operativa"
                
        elif self.estado == "parando":
            self.vel_angular -= self.vel_angular/self.tiempo_parada
            if self.vel_angular <= 50:
                self.estado = "detenida"
                
        return self.curva_qh(flujo_actual) * (self.vel_angular/1500)
    

class ValvulaNoRetorno:
    def __init__(self, node_id, tiempo_cierre=1.0):
        self.node_id = node_id
        self.factor_apertura = 1.0
        self.tiempo_cierre = tiempo_cierre
        self.estado = "abierta"
        
    def actualizar(self, flujo, dt):
        if flujo < 0 and self.estado == "abierta":
            self.estado = "cerrando"
            self.tiempo_transcurrido = 0.0
            
        if self.estado == "cerrando":
            self.tiempo_transcurrido += dt
            self.factor_apertura = 1 - (self.tiempo_transcurrido/self.tiempo_cierre)
            if self.tiempo_transcurrido >= self.tiempo_cierre:
                self.estado = "cerrada"
                self.factor_apertura = 0.0
                
    def factor_apertura(self):
        return 1.0 if self.abierta else 0.0


def simular_red_epanet(archivo_inp, duracion=3600):
    try:

        ret = et.ENopen(archivo_inp, "reporte.txt", "")
        
        et.ENopenH()
        et.ENinitH(0) 
        
        tiempos = []
        presiones = {}
        flujos = {}
        
        while True:
            errcode, t = et.ENrunH() 
            tiempos.append(t)
            
            errcode, n_nodes = et.ENgetcount(et.EN_NODECOUNT)
            for idx in range(1, n_nodes + 1):
                errcode, presion = et.ENgetnodevalue(idx, et.EN_PRESSURE)
                presiones.setdefault(idx, []).append(presion)
            
            errcode, n_links = et.ENgetcount(et.EN_LINKCOUNT)
            for idx in range(1, n_links + 1):
                errcode, flujo = et.ENgetlinkvalue(idx, et.EN_FLOW)
                flujos.setdefault(idx, []).append(flujo)
            
            if t >= duracion:
                break
                
            et.ENnextH()
        
        et.ENcloseH()
        et.ENclose()
        
        return {
            'tiempos': np.array(tiempos),
            'presiones': presiones,
            'flujos': flujos
        }
        
    except Exception as e:
        et.ENclose()
        raise e

        
class ProblemaHidraulico(Problem):
    def __init__(self, red_epanet):
        super().__init__(n_var=3, n_obj=2, n_constr=1, xl=np.array([0.1, 0.5, 1.0]), xu=np.array([2.0, 5.0, 10.0]))
        self.red = red_epanet
        
    def _evaluate(self, X, out, *args, **kwargs):
        F = np.zeros((X.shape[0], 2))
        G = np.zeros((X.shape[0], 1))
        
        for i in range(X.shape[0]):
            diam, v_ini, t_cierre = X[i]
            
            t, Hup, Hdown = metodo_caracteristicas_simple(
                L=self.red['L'], D=diam, v_inicial=v_ini, 
                t_cierre_valvula=t_cierre
            )
            
            costo = diam * self.red['L'] * 1500
            sobrepresion = np.max(Hdown) - self.red['H_down']
            
            F[i, 0] = costo
            F[i, 1] = sobrepresion
            
            G[i, 0] = 20 - np.min(Hdown)
            
        out["F"] = F
        out["G"] = G

def worker_simulacion(archivo, queue):
    try:
        data = leer_epanet_principal(archivo)
        if 'error' in data:
            queue.put(('error', data['error']))
            return
        
        queue.put(('parametros', data))
        
        resultados = simular_red_epanet(archivo)
        queue.put(('resultados', resultados))
        
    except Exception as e:
        queue.put(('error', str(e)))

def leer_epanet_principal(filename):
    params = {
        'rho': 1000.0,
        'c': 1200.0,
        'L': 100.0,
        'D': 0.5,
        'P_inicial': 300
    }
    
    with open(filename, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        for line in lines:
            if line.strip().startswith(";"):
                if match := re.search(r'rho\s*=\s*([\d\.]+)', line, re.IGNORECASE):
                    params['rho'] = float(match.group(1))
                if match := re.search(r'c\s*=\s*([\d\.]+)', line, re.IGNORECASE):
                    params['c'] = float(match.group(1))
    
    try:
        err = et.ENopen(filename, "report.txt", "")
        errcode, n_links = et.ENgetcount(et.EN_LINKCOUNT)
        selected_link_idx = None

        for idx in range(1, n_links + 1):
            errcode, link_type = et.ENgetlinktype(idx)
            if errcode == 0:
                errcode, length = et.ENgetlinkvalue(idx, et.EN_LENGTH)
                errcode, diameter = et.ENgetlinkvalue(idx, et.EN_DIAMETER)
                if errcode == 0:
                    params['L'] = length
                    params['D'] = diameter
                    selected_link_idx = idx
                    break

        if selected_link_idx is not None:
            errcode, node1, node2 = et.ENgetlinknodes(selected_link_idx)
            errcode, elev1 = et.ENgetnodevalue(node1, et.EN_ELEVATION)
            errcode, elev2 = et.ENgetnodevalue(node2, et.EN_ELEVATION)
            params['H_up'] = elev1 + 10.0
            params['H_down'] = elev2 + 5.0

        et.ENclose()
        return params
        
    except Exception as e:
        et.ENclose()
        return {'error': f"Error EPANET: {str(e)}"}
    

def metodo_caracteristicas_avanzado(rho, c, L, D,
                                    v_inicial=1.0,
                                    t_cierre_valvula=1.0,
                                    t_pump_start=1.0, 
                                    t_pump_stop=2.0,
                                    scenario="sin_bomba",
                                    flujo_bomba=None,
                                    presion_bomba=None,
                                    t_total=5.0,
                                    dt=0.01,
                                    P_inicial=300000,  
                                    bombas=[], 
                                    valvulas=[]):
    
    g = 9.81
    H0 = P_inicial / (rho * g)
    
    A = np.pi*(D**2)/4.0
    nsteps = int(t_total/dt) + 1
    t_array = np.linspace(0, t_total, nsteps)
    
    Hup = np.zeros(nsteps)
    Hdown = np.zeros(nsteps)
    Qline = np.zeros(nsteps)
    
    Hup[0] = H0
    Hdown[0] = H0
    Qline[0] = v_inicial * A

    f = 0.02
    Beta = f * L / (2*g*D)
    alpha = c/(g*A)

    for n in range(1, nsteps):

        t_now = t_array[n]
        Q_old = Qline[n-1]

        for bomba in bombas:
            H_bomba = bomba.actualizar_estado(t_now, Q_old)
            Hup[n] += H_bomba

        if scenario in ["sin_bomba", "cierre_valvula"]:

            if scenario == "sin_bomba":
                valve_factor = 1.0 - (t_now/t_cierre_valvula) if t_now < t_cierre_valvula else 0.0

            else:
                valve_factor = 1.0 if t_now < t_cierre_valvula else 0.0
            Hup[n] = H0
            c_plus = Hup[n] - Beta * Q_old * abs(Q_old) + alpha * Q_old
            c_minus = Hdown[n-1] + Beta * Q_old * abs(Q_old) - alpha * Q_old
            Qc = (c_plus - c_minus) / (2*Beta + 1e-12)
            Q_new = valve_factor * Qc

        elif scenario == "con_bomba":
            if t_now < t_pump_start:
                pump_factor = 0.0
            elif t_now < t_pump_stop:
                pump_factor = (t_now - t_pump_start) / (t_pump_stop - t_pump_start)
            else:
                pump_factor = 1.0
            pump_head = pump_factor * (presion_bomba/(rho*g)) if presion_bomba is not None else 0.0
            Hup[n] = H0 + pump_head
            c_plus = Hup[n] - Beta * Q_old * abs(Q_old) + alpha * Q_old
            c_minus = Hdown[n-1] + Beta * Q_old * abs(Q_old) - alpha * Q_old
            Q_new = (c_plus - c_minus) / (2*Beta + 1e-12)
        else:
            Hup[n] = H0
            c_plus = Hup[n] - Beta * Q_old * abs(Q_old) + alpha * Q_old
            c_minus = Hdown[n-1] + Beta * Q_old * abs(Q_old) - alpha * Q_old
            Q_new = (c_plus - c_minus) / (2*Beta + 1e-12)

        for valvula in valvulas:
            valvula.actualizar(Q_old, dt)
            factor = valvula.factor_apertura
            Q_new *= factor

            Hup[n] = H0

    Pup = np.array(Hup) * rho * g
    Pdown = np.array(Hdown) * rho * g

    return t_array, Pup, Pdown, Qline


class AppGolpeAriete(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("EPANET-PYTHON con MOC - Avanzado")
        self.geometry("1380x900") 
        self.config(bg="#f0f8ff")
        self.resizable(False, False)
        self.progress_window = None  
        self.simulation_thread = None  
        self.simulation_process = None
        self.queue = mp.Queue()
        self.water_hammer_analyzer = None

        self.parametros = {
            'rho': {"label": "Densidad (kg/m³):", "entry": None},
            'c': {"label": "Velocidad de onda (m/s):", "entry": None},
            'L': {"label": "Longitud (m):", "entry": None},
            'D': {"label": "Diámetro (m):", "entry": None},
            'v_inicial': {"label": "Vel. inicial (m/s):", "entry": None},
            'tiempo_cierre_valvula': {"label": "Tiempo cierre válvula (s):", "entry": None},
            't_total': {"label": "Tiempo total sim (s):", "entry": None},
            'dt': {"label": "Paso de tiempo dt (s):", "entry": None},
            'P_inicial': {"label": "Presión inicial (kPa):", "entry": None}
        }

        self.flujo_bomba = None
        self.presion_bomba = None

        self.scenario = tk.StringVar(value="sin_bomba")

        self.crear_interfaz()
        self.graph_counter = 0

        self.P_inicial = 300

        self.bombas = []
        self.valvulas = []
        self.datos_epanet = None

    def cargar_archivo(self):
        archivo = filedialog.askopenfilename(
            title="Selecciona un archivo EPANET (.inp)",
            filetypes=[("Archivos EPANET", "*.inp"), ("Todos los archivos", "*.*")]
        )
        if not archivo:
            return

        self.water_hammer_analyzer = WaterHammerAnalyzer(archivo)
        self.water_hammer_analyzer.identify_critical_nodes()

        self.mostrar_progreso()
        self.simulation_process = mp.Process(
            target=worker_simulacion,
            args=(archivo, self.queue),
            daemon=True
        )
        self.simulation_process.start()
        self.verificar_estado_proceso()

        self.dibujar_red_epanet()

    def ejecutar_simulacion_epanet(self, archivo, queue):
        try:
            data = leer_epanet_principal(archivo)
            queue.put(('parametros', data))
            
            resultados = simular_red_epanet(archivo)
            queue.put(('resultados', resultados))
            
        except Exception as e:
            queue.put(('error', str(e)))

    def verificar_estado_proceso(self):
        try:
            while True:
                msg = self.queue.get_nowait()
                if msg[0] == 'parametros':
                    self.actualizar_parametros(msg[1])
                elif msg[0] == 'resultados':
                    self.datos_epanet = msg[1]
                    messagebox.showinfo("Éxito", "Simulación completada exitosamente")
                elif msg[0] == 'error':
                    messagebox.showerror("Error", msg[1])
        except Empty:
            if self.simulation_process.is_alive():
                self.after(100, self.verificar_estado_proceso)
            else:
                self.ocultar_progreso()

    def actualizar_parametros(self, data):
        if 'error' in data:
            messagebox.showerror("Error", data['error'])
            return
            
        for key in ['rho', 'c', 'L', 'D']:
            self.parametros[key]["entry"].delete(0, tk.END)
            self.parametros[key]["entry"].insert(0, str(data[key]))
        
        self.H_up = data['H_up']
        self.H_down = data['H_down']

    def mostrar_progreso(self):
        self.progress_window = tk.Toplevel(self)
        self.progress_window.title("Procesando...")
        
        frame = tk.Frame(self.progress_window, padx=20, pady=20)
        frame.pack()
        
        tk.Label(frame, text="Simulando red EPANET...", font=("Arial", 12)).pack(pady=10)
        
        self.progress_bar = ttk.Progressbar(
            frame, 
            mode='indeterminate',
            length=300
        )
        self.progress_bar.pack(pady=10)
        self.progress_bar.start(15)
        
        self.progress_window.geometry("+%d+%d" % (
            self.winfo_x() + self.winfo_width()/2 - 150,
            self.winfo_y() + self.winfo_height()/2 - 50
        ))
        self.progress_window.grab_set()

    def ocultar_progreso(self):
        if self.progress_window:
            self.progress_bar.stop()
            self.progress_window.destroy()
            self.progress_window = None

    def verificar_estado_hilo(self):
        if self.simulation_thread.is_alive():
            self.after(100, self.verificar_estado_hilo)
        else:
            self.after(0, self.fin_carga_archivo)

    def fin_carga_archivo(self):
        self.ocultar_progreso()
        messagebox.showinfo("Archivo cargado", "Simulación EPANET completada exitosamente")

    def validar_con_epanet(self):
        if not self.datos_epanet:
            messagebox.showerror("Error", "Primero carga un archivo EPANET.")
            return

        fig = plt.figure(figsize=(10, 6))
        ax1 = fig.add_subplot(211)
        ax2 = fig.add_subplot(212)
        
        params = {key: float(self.parametros[key]["entry"].get()) for key in ['rho', 'c', 'L', 'D']}
        t_moc, Hup_moc, Hdown_moc, Qline = metodo_caracteristicas_simple(**params)
        
        nodo_objetivo = 1 
        ax1.plot(self.datos_epanet['tiempos'], self.datos_epanet['presiones'][nodo_objetivo], 
                label='EPANET', linestyle='--')
        ax1.plot(t_moc, Hdown_moc, label='MOC')
        ax1.set_title("Comparación de Presiones")
        
        tuberia_objetivo = 1 
        ax2.plot(self.datos_epanet['tiempos'], self.datos_epanet['flujos'][tuberia_objetivo],
                label='EPANET', linestyle='--')
        ax2.plot(t_moc, [q * 1000 for q in Qline], label='MOC') 
        ax2.set_title("Comparación de Flujos")
        
        plt.legend()
        self.mostrar_grafico(fig)

    def optimizacion_avanzada(self):
        optimizador = OptimizadorHidraulico(self.water_hammer_analyzer.network)
        resultados = optimizador.optimizar_nsga2()
        self.mostrar_resultados_optimizacion(resultados)
        try:
            datos_red = {
                'L': float(self.parametros["L"]["entry"].get()),
                'H_down': self.H_down
            }
            
            problema = ProblemaHidraulico(datos_red)
            algoritmo = NSGA2(pop_size=50)
            res = minimize(problema, algoritmo, ('n_gen', 100))
            
            soluciones = res.X
            objetivos = res.F
            
            fig = plt.figure()
            ax = fig.add_subplot(111)
            ax.scatter(objetivos[:, 0], objetivos[:, 1], c='red')
            ax.set_xlabel("Costo")
            ax.set_ylabel("Sobrepresión")
            self.mostrar_grafico(fig)
            
        except Exception as e:
            messagebox.showerror("Error", f"Error en optimización: {str(e)}")

    def analizar_golpe_ariete(self):
        if not self.water_hammer_analyzer:
            messagebox.showerror("Error", "Cargar archivo .inp primero")
            return
            
        close_time = float(self.parametros["tiempo_cierre_valvula"]["entry"].get())
        results = self.water_hammer_analyzer.run_analysis(close_time)
        
        doc = Document()
        doc.add_heading('Análisis de Golpe de Ariete', 0)
        
        for pipe_id, data in results.items():
            doc.add_heading(f'Tubería {pipe_id}', level=1)
            doc.add_paragraph(f"Sobrepresión Simulada: {data['simulated']/1e6:.2f} MPa")
            doc.add_paragraph(f"Joukowsky: {data['joukowsky']/1e6:.2f} MPa")
            doc.add_paragraph(f"Relación Sim/Jouk: {data['simulated']/data['joukowsky']:.2f}")
            
            fig = plt.figure()
            plt.plot(data['data'][0], data['data'][1])
            plt.title(f"Presión en Tubería {pipe_id}")
            img_stream = BytesIO()
            fig.savefig(img_stream)
            doc.add_picture(img_stream, width=Inches(5))
            
        doc.save("Informe_Golpe_Ariete.docx")
        messagebox.showinfo("Listo", "Reporte generado en Informe_Golpe_Ariete.docx")

    def cargar_caso_peruano(self):
        datos = {
            'rho': 1000,
            'c': 1200,
            'L': 2500,  
            'D': 0.6,   
            'H_up': 1500, 
            'H_down': 1450
        }
        
        for key in ['rho', 'c', 'L', 'D']:
            self.parametros[key]["entry"].delete(0, tk.END)
            self.parametros[key]["entry"].insert(0, str(datos[key]))
        
        self.H_up = datos['H_up']
        self.H_down = datos['H_down']
        messagebox.showinfo("Caso Peruano", "Parámetros de la Sierra cargados.")

    def mostrar_grafico(self, fig):
        ventana_grafico = tk.Toplevel(self)
        ventana_grafico.title("Resultados Avanzados")
        canvas = FigureCanvasTkAgg(fig, master=ventana_grafico)
        canvas.draw()
        canvas.get_tk_widget().pack()

    def crear_interfaz(self):
        frame = tk.Frame(self, bg="#ffffff", padx=30, pady=30, relief="solid", bd=2)
        frame.place(relx=0.05, rely=0.5, anchor="w")

        titulo = tk.Label(frame, text="Parámetros", font=("Arial", 20, "bold"), bg="#ffffff", fg="#00796b")
        titulo.grid(row=0, column=0, columnspan=2, pady=(0, 20))

        row_index = 1
        for key, info in self.parametros.items():
            label = tk.Label(frame, text=info["label"], bg="#ffffff", font=("Arial", 10, "italic"), anchor="w")
            label.grid(row=row_index, column=0, padx=10, pady=5, sticky="w")

            entry = tk.Entry(frame, font=("Arial", 10), width=20, relief="solid", bd=1,
                             highlightthickness=2, highlightbackground="#00796b")
            entry.grid(row=row_index, column=1, padx=10, pady=5)
            self.parametros[key]["entry"] = entry
            row_index += 1

        escenario_label = tk.Label(frame, text="Escenario:", bg="#ffffff", font=("Arial", 10, "italic"))
        escenario_label.grid(row=row_index, column=0, padx=10, pady=5, sticky="w")
        rb1 = tk.Radiobutton(frame, text="Base (Sin Bomba)", variable=self.scenario, value="sin_bomba", bg="#ffffff")
        rb2 = tk.Radiobutton(frame, text="Con Bomba (Arranque/Parada)", variable=self.scenario, value="con_bomba", bg="#ffffff")
        rb3 = tk.Radiobutton(frame, text="Cierre Repentino Válvula", variable=self.scenario, value="cierre_valvula", bg="#ffffff")
        rb4 = tk.Radiobutton(frame, text="Modelo Inercial Rígido", variable=self.scenario, value="rigido", bg="#ffffff")
        rb1.grid(row=row_index, column=1, padx=10, pady=2, sticky="w")
        rb2.grid(row=row_index+1, column=1, padx=10, pady=2, sticky="w")
        rb3.grid(row=row_index+2, column=1, padx=10, pady=2, sticky="w")
        rb4.grid(row=row_index+3, column=1, padx=10, pady=2, sticky="w")
        row_index += 4

        button_cargar = tk.Button(
            frame, text="Cargar archivo EPANET", command=self.cargar_archivo,
            font=("Arial", 12), bg="#00796b", fg="white", relief="flat", padx=15, pady=10
        )
        button_cargar.grid(row=row_index, column=0, columnspan=2, pady=(20, 10))
        row_index += 1

        button_calcular = tk.Button(
            frame, text="Calcular y Graficar (MOC)", command=self.calcular,
            font=("Arial", 12), bg="#00796b", fg="white", relief="flat", padx=15, pady=10
        )
        button_calcular.grid(row=row_index, column=0, columnspan=2, pady=10)
        row_index += 1

        button_analizar = tk.Button(
            self, text="Realizar Análisis", command=self.realizar_analisis,
            font=("Arial", 12), bg="#00796b", fg="white", relief="flat", padx=15, pady=10
        )
        button_analizar.place(relx=0.1, rely=0.1, anchor="center")

        button_insertar_bomba = tk.Button(
            self, text="Insertar Bomba", command=self.insertar_bomba,
            font=("Arial", 12), bg="#00796b", fg="white", relief="flat", padx=15, pady=10
        )
        button_insertar_bomba.place(relx=0.25, rely=0.1, anchor="center")

        button_opt = tk.Button(
            self, text="Optimizar Sistema", command=self.optimizar_sistema,
            font=("Arial", 12), bg="#FF5722", fg="white", relief="flat", padx=15, pady=10
        )
        button_opt.place(relx=0.4, rely=0.1, anchor="center")

        button_validar = tk.Button(
            self, text="Validar vs EPANET", command=self.validar_con_epanet,
            font=("Arial", 12), bg="#4CAF50", fg="white", relief="flat", padx=15, pady=10
        )
        button_validar.place(relx=0.55, rely=0.1, anchor="center")

        button_opt_avanzada = tk.Button(
            self, text="Optimización Avanzada", command=self.optimizacion_avanzada,
            font=("Arial", 12), bg="#FF5722", fg="white", relief="flat", padx=15, pady=10
        )
        button_opt_avanzada.place(relx=0.7, rely=0.1, anchor="center")

        button_caso_peru = tk.Button(
            self, text="Cargar Caso Peruano", command=self.cargar_caso_peruano,
            font=("Arial", 12), bg="#2196F3", fg="white", relief="flat", padx=15, pady=10
        )
        button_caso_peru.place(relx=0.85, rely=0.1, anchor="center")

        self.canvas_frame = tk.Frame(self)
        self.canvas_frame.place(relx=0.7, rely=0.5, anchor="center")

        self.canvas = tk.Canvas(self.canvas_frame, bg="white", width=800, height=600, scrollregion=(0, 0, 5000, 5000))

        self.scrollbar = tk.Scrollbar(self.canvas_frame, orient="vertical", command=self.canvas.yview)
        self.scrollbar.pack(side="right", fill="y")
        self.canvas.config(yscrollcommand=self.scrollbar.set)
        self.canvas.pack()

        self.graph_frame = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.graph_frame, anchor="nw")

    def insertar_bomba(self):
        ventana_bomba = tk.Toplevel(self)
        ventana_bomba.title("Parámetros de la Bomba")
        ventana_bomba.geometry("300x200")
        ventana_bomba.config(bg="#f0f8ff")
        
        parametros_bomba = {
            'flujo': {"label": "Flujo (m³/h):", "entry": None},
            'presion': {"label": "Presión (bar):", "entry": None},
            'potencia': {"label": "Potencia (kW):", "entry": None}
        }

        row_index = 1
        for key, info in parametros_bomba.items():
            label = tk.Label(ventana_bomba, text=info["label"], bg="#ffffff",
                             font=("Arial", 10), anchor="w")
            label.grid(row=row_index, column=0, padx=10, pady=5, sticky="w")
            entry = tk.Entry(ventana_bomba, font=("Arial", 10), width=20, relief="solid",
                             bd=1, highlightthickness=2, highlightbackground="#00796b")
            entry.grid(row=row_index, column=1, padx=10, pady=5)
            parametros_bomba[key]["entry"] = entry
            row_index += 1

        button_guardar = tk.Button(
            ventana_bomba, text="Guardar Parámetros",
            command=lambda: self.guardar_parametros_bomba(parametros_bomba),
            font=("Arial", 12), bg="#00796b", fg="white", relief="flat", padx=15, pady=10
        )
        button_guardar.grid(row=row_index, column=0, columnspan=2, pady=10)

    def guardar_parametros_bomba(self, parametros_bomba):
        try:
            flujo = float(parametros_bomba["flujo"]["entry"].get())
            presion = float(parametros_bomba["presion"]["entry"].get())
            potencia = float(parametros_bomba["potencia"]["entry"].get())
            messagebox.showinfo(
                "Parámetros Guardados",
                f"Flujo: {flujo} m³/h\nPresión: {presion} bar\nPotencia: {potencia} kW"
            )
            self.flujo_bomba = flujo
            self.presion_bomba = presion
        except ValueError:
            messagebox.showerror(
                "Error",
                "Por favor, ingresa valores numéricos válidos para la bomba."
            )

    def mostrar_resultados_por_tuberia(self, results):
        ventana_resultados = tk.Toplevel(self)
        ventana_resultados.title("Resultados por Tubería")
        ventana_resultados.geometry("1200x800")
        
        notebook = ttk.Notebook(ventana_resultados)
        notebook.pack(fill='both', expand=True)

        for pipe_id, data in results.items():
            frame = tk.Frame(notebook)
            notebook.add(frame, text=f"Tubería {pipe_id}")
            
            fig = plt.figure(figsize=(10, 6))
            ax = fig.add_subplot(111)
            ax.plot(data['data'][0], data['data'][1])
            ax.set_title(f"Presión en {pipe_id}")
            ax.set_xlabel('Tiempo (s)')
            ax.set_ylabel('Altura Piezométrica (m)')
            
            canvas = FigureCanvasTkAgg(fig, frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill='both', expand=True)
            
            params_text = f"Longitud: {data['params']['L']:.2f} m\nDiámetro: {data['params']['D']:.2f} m\nVelocidad inicial: {data['params']['v0']:.2f} m/s"
            label_params = tk.Label(frame, text=params_text, justify='left')
            label_params.pack()

    def editar_parametros_tuberia(self, pipe_id):
        data = self.water_hammer_analyzer.pipes_data[pipe_id]
        
        ventana_edicion = tk.Toplevel(self)
        ventana_edicion.title(f"Editar {pipe_id}")
        
        entries = {}
        for i, (key, value) in enumerate(data.items()):
            if key in ['L', 'D', 'v0']:
                tk.Label(ventana_edicion, text=key).grid(row=i, column=0)
                entries[key] = tk.Entry(ventana_edicion)
                entries[key].insert(0, str(value))
                entries[key].grid(row=i, column=1)
        
        tk.Button(ventana_edicion, text="Guardar", 
                command=lambda: self.guardar_cambios(pipe_id, entries)).grid(row=len(entries)+1, columnspan=2)

    def guardar_cambios(self, pipe_id, entries):
        new_data = {
            'L': float(entries['L'].get()),
            'D': float(entries['D'].get()),
            'v0': float(entries['v0'].get())
        }
        self.water_hammer_analyzer.pipes_data[pipe_id].update(new_data)
        messagebox.showinfo("Éxito", "Parámetros actualizados")


    def calcular(self):
        rho = float(self.parametros["rho"]["entry"].get())
        c = float(self.parametros["c"]["entry"].get())
        L = float(self.parametros["L"]["entry"].get())
        D = float(self.parametros["D"]["entry"].get())
        v_inicial = float(self.parametros["v_inicial"]["entry"].get())
        t_cierre = float(self.parametros["tiempo_cierre_valvula"]["entry"].get())
        t_total = float(self.parametros["t_total"]["entry"].get()) if self.parametros["t_total"]["entry"].get() else 5.0
        dt = float(self.parametros["dt"]["entry"].get()) if self.parametros["dt"]["entry"].get() else 0.01

        escenario = self.scenario.get()
        t_pump_start = 1.0
        t_pump_stop = 2.0

        if escenario == "rigido":
            t, P, Q = modelo_inercial_rigido(
                rho=rho,
                L=L,
                D=D,
                v_inicial=v_inicial,
                t_cierre_valvula=t_cierre,
                t_total=t_total,
                dt=dt,
                P_inicial=self.P_inicial * 1000
            )
            
            Qline = Q
            t1 = t
            Pup1 = P
            Pdown1 = P 

        else:
            t1, Pup1, Pdown1, Qline = metodo_caracteristicas_avanzado(
                rho, c, L, D,
                v_inicial=v_inicial,
                t_cierre_valvula=t_cierre,
                t_pump_start=t_pump_start, 
                t_pump_stop=t_pump_stop,
                scenario=escenario,
                flujo_bomba=self.flujo_bomba,
                presion_bomba=self.presion_bomba,
                t_total=t_total,
                dt=dt,
                P_inicial=self.P_inicial * 1000  
            )

        g = 9.81

        for widget in self.graph_frame.winfo_children():
            widget.destroy()

        fig1, ax1 = plt.subplots(figsize=(8, 5))
        
        ax1.plot(t1, np.array(Pup1)/1000, label="Upstream", linewidth=2)
        ax1.plot(t1, np.array(Pdown1)/1000, label="Downstream", linewidth=2)
        ax1.set_ylabel('Presión (kPa)')
        
        ax1.set_title(f"Simulación Hidráulica - {escenario}", fontsize=14, weight='bold')
        ax1.set_xlabel('Tiempo (s)', fontsize=12, labelpad=10)
        ax1.set_ylabel('Altura Piezométrica (m)', fontsize=12, labelpad=10)
        ax1.grid(True, linestyle='--', alpha=0.7)
        ax1.legend(loc='upper right', frameon=True, shadow=True)
        
        textstr = '\n'.join((
            f'ΔP máximo: {np.max(Pdown1)/1000 - self.P_inicial:.2f} kPa',
            f'v₀: {v_inicial} m/s',
            f'L: {L} m',
            f'D: {D} m',
            f'c: {c} m/s'))
        ax1.text(0.95, 0.15, textstr, transform=ax1.transAxes,
                fontsize=10, verticalalignment='top', 
                horizontalalignment='right',
                bbox=dict(boxstyle='round', 
                        facecolor='white', 
                        alpha=0.8,
                        edgecolor='#00796b'))

        canvas1 = FigureCanvasTkAgg(fig1, self.graph_frame)
        canvas1.draw()
        canvas1.get_tk_widget().pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

        fig_tabla = self.generar_tabla_resultados(t1, Pup1, Pdown1, Qline)

        canvas_tabla = FigureCanvasTkAgg(fig_tabla, self.graph_frame)
        canvas_tabla.draw()
        canvas_tabla.get_tk_widget().pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

        ventana_resultados = tk.Toplevel(self)
        ventana_resultados.title("Resultados de Simulación MOC")
        ventana_resultados.geometry("1000x800")
        
        graph_container = tk.Frame(ventana_resultados)
        graph_container.pack(fill=tk.BOTH, expand=True)
        
        canvas1 = FigureCanvasTkAgg(fig1, master=graph_container)
        canvas1.draw()
        canvas1.get_tk_widget().pack(pady=10, padx=10, fill=tk.BOTH, expand=True)
        
        canvas_tabla = FigureCanvasTkAgg(fig_tabla, master=graph_container)
        canvas_tabla.draw()
        canvas_tabla.get_tk_widget().pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

        if not hasattr(self, 'figures'):
            self.figures = []
        self.figures.extend([fig1, fig_tabla])

        self.canvas.config(scrollregion=self.canvas.bbox("all"))
        
        messagebox.showinfo("Resultado", "Simulación completada con éxito.\nRevisa los gráficos y tablas para los resultados detallados.")
        
        if hasattr(self, 'figures'):
            for fig in self.figures:
                plt.close(fig)
        self.figures = [fig1, fig_tabla]
        
        messagebox.showinfo("Resultado", "Gráficos generados en ventana emergente.")

        self.water_hammer_analyzer.load_pipe_parameters()
        close_time = float(self.parametros["tiempo_cierre_valvula"]["entry"].get())
        results = self.water_hammer_analyzer.run_analysis(close_time)
        self.mostrar_resultados_por_tuberia(results)

    def generar_tabla_resultados(self, t, Hup, Hdown, Q):
        df = pd.DataFrame({
            'Tiempo (s)': t,
            'Altura Upstream (m)': Hup,
            'Altura Downstream (m)': Hdown,
            'Caudal (m³/s)': Q
        })
        
        stats = pd.DataFrame({
            'Máximo': [df['Altura Upstream (m)'].max(), 
                    df['Altura Downstream (m)'].max(), 
                    df['Caudal (m³/s)'].max()],
            'Mínimo': [df['Altura Upstream (m)'].min(), 
                    df['Altura Downstream (m)'].min(), 
                    df['Caudal (m³/s)'].min()],
            'Promedio': [df['Altura Upstream (m)'].mean(), 
                        df['Altura Downstream (m)'].mean(), 
                        df['Caudal (m³/s)'].mean()]
        }, index=['Upstream', 'Downstream', 'Caudal'])
        
        fig, ax = plt.subplots(figsize=(8, 3))
        ax.axis('tight')
        ax.axis('off')
        tabla = ax.table(cellText=stats.values.round(2),
                        colLabels=stats.columns,
                        rowLabels=stats.index,
                        cellLoc='center',
                        loc='center',
                        bbox=[0, 0, 1, 1])
        
        tabla.auto_set_font_size(False)
        tabla.set_fontsize(10)
        tabla.scale(1.2, 1.2)
        
        return fig

    def realizar_analisis(self):
        try:
            rho = float(self.parametros["rho"]["entry"].get())
            c = float(self.parametros["c"]["entry"].get())
            L = float(self.parametros["L"]["entry"].get())
            D = float(self.parametros["D"]["entry"].get())
            v_inicial = float(self.parametros["v_inicial"]["entry"].get())
            t_cierre = float(self.parametros["tiempo_cierre_valvula"]["entry"].get())
            
            t_array, hup, hdown, _ = metodo_caracteristicas_avanzado(
                rho, c, L, D,
                v_inicial=v_inicial,
                t_cierre_valvula=t_cierre,
                scenario=self.scenario.get(),
                t_total=5.0,
                dt=0.01,
                H_up=self.H_up,
                H_down=self.H_down
            )
            
            g = 9.81
            p_up = rho * g * hup  
            p_down = rho * g * hdown
            max_p = max(np.max(p_up), np.max(p_down))
            base_p = rho * g * self.H_down
            delta_p = max_p - base_p
            analisis = self.generar_analisis(delta_p, rho, c, L, D, v_inicial)
            self.mostrar_analisis(analisis)
            
        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un error: {e}")

    def generar_analisis(self, delta_P, rho, c, L, D, v_inicial):

        analisis = ""

        if self.scenario.get() == "rigido":
            analisis += "**Modelo Utilizado**: Dinámico Inercial Rígido\n"
            analisis += "- Considera inercia del fluido\n"
            analisis += "- Ignora efectos elásticos y onda de presión\n"
            analisis += "- Apropiado para eventos lentos (ej: arranque bombas)\n\n"
        else:
            analisis += "**Modelo Utilizado**: Método de las Características (Elástico)\n\n"

        analisis = "Análisis Detallado del Golpe de Ariete\n"
        analisis += "=====================================\n\n"
        
        analisis += "**Parámetros de Simulación:**\n"
        analisis += f"- Densidad del fluido (ρ): {rho} kg/m³\n"
        analisis += f"- Velocidad de onda de presión (c): {c} m/s\n"
        analisis += f"- Longitud de tubería (L): {L} m\n"
        analisis += f"- Diámetro interno (D): {D} m\n"
        analisis += f"- Velocidad inicial del flujo (v₀): {v_inicial} m/s\n"
        analisis += f"- Presión inicial: {self.P_inicial} kPa\n"
        
        A = np.pi*(D**2)/4
        Q0 = v_inicial * A
        tiempo_reflexion = 2*L/c
        J = (c * v_inicial)/9.81
        
        J_presion = rho * c * v_inicial
        analisis += f"- Sobrepresión teórica (Joukowsky): {J_presion/1000:.2f} kPa\n"

        analisis += "**Cálculos Fundamentales:**\n"
        analisis += f"- Área transversal de la tubería: {A:.4f} m²\n"
        analisis += f"- Caudal inicial (Q₀): {Q0:.4f} m³/s\n"
        analisis += f"- Tiempo de reflexión de onda: {tiempo_reflexion:.2f} s\n"
        analisis += f"- Sobrepresión teórica (Joukowsky): {J:.2f} m\n\n"
        
        analisis += "**Resultados de la Simulacion:**\n"
        analisis += f"- Sobrepresión máxima registrada: {delta_P/1000:.2f} kPa\n"
        analisis += f"- Relación Joukowsky/Resultado: {(J*9.81*rho)/(delta_P):.2f}\n"
        
        analisis += "\n**Análisis Comparativo:**\n"
        if abs((J*9.81*rho) - delta_P) < 0.1*delta_P:
            analisis += "La sobrepresión se alinea con la teoría de Joukowsky (variación <10%)\n"
        else:
            analisis += "Se detectan discrepancias significativas respecto a la teoría base\n"
        
        analisis += "\n**Recomendaciones Técnicas:**\n"
        if D < 0.3:
            analisis += "- Considerar tuberías de mayor diámetro para reducir velocidades\n"
        if c < 800:
            analisis += "- Evaluar material de tubería (módulo de elasticidad)\n"
        
        analisis += "\n**Efectos del Escenario Simulado:**\n"
        if self.scenario.get() == "con_bomba":
            analisis += "- La operación de bombas introduce variaciones de presión cíclicas\n"
            analisis += f"- Potencia de bomba: {self.presion_bomba} bar\n"
            analisis += f"- Flujo de bomba: {self.flujo_bomba} m³/h\n"
        elif self.scenario.get() == "cierre_valvula":
            analisis += "- El cierre repentino genera ondas de presión de alta frecuencia\n"
        
        return analisis

    def mostrar_analisis(self, analisis):
        analisis_window = tk.Toplevel(self)
        analisis_window.title("Análisis del Golpe de Ariete")
        analisis_window.geometry("600x400")
        analisis_window.resizable(False, False)
        analisis_window.config(bg="#ffffff")
        texto_analisis = tk.Text(analisis_window, wrap=tk.WORD, width=70, height=15, font=("Arial", 10))
        texto_analisis.insert(tk.END, analisis)
        texto_analisis.config(state=tk.DISABLED)
        texto_analisis.pack(padx=20, pady=20)
        button_guardar_word = tk.Button(
            analisis_window, text="Guardar en Word",
            command=lambda: self.guardar_en_word(analisis),
            font=("Arial", 12), bg="#00796b", fg="white", relief="flat"
        )
        button_guardar_word.pack(pady=5)
        button_guardar_pdf = tk.Button(
            analisis_window, text="Guardar en PDF",
            command=lambda: self.guardar_en_pdf(analisis),
            font=("Arial", 12), bg="#00796b", fg="white", relief="flat"
        )
        button_guardar_pdf.pack(pady=5)
        boton_cerrar = tk.Button(
            analisis_window, text="Cerrar",
            command=analisis_window.destroy,
            font=("Arial", 12), bg="#FF5722", fg="white", relief="flat"
        )
        boton_cerrar.pack(pady=10)

    def guardar_en_word(self, analisis):
        doc = Document()
        doc.add_heading('Análisis del Golpe de Ariete (MOC Avanzado)', 0)
        doc.add_paragraph(analisis)
        archivo_guardar = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documentos Word", "*.docx")]
        )
        if archivo_guardar:
            doc.save(archivo_guardar)
            messagebox.showinfo("Guardado", f"El análisis se guardó en: {archivo_guardar}")

    def guardar_en_pdf(self, analisis):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Análisis del Golpe de Ariete (MOC Avanzado)", ln=True, align="C")
        pdf.ln(10)
        pdf.multi_cell(0, 10, analisis)
        archivo_guardar = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("Documentos PDF", "*.pdf")]
        )
        if archivo_guardar:
            pdf.output(archivo_guardar)
            messagebox.showinfo("Guardado", f"El análisis se guardó en: {archivo_guardar}")

    def optimizar_sistema(self):
        messagebox.showinfo(
            "Optimización",
            "Se buscará la combinación (v_inicial, t_cierre) que minimice la sobrepresión máxima."
        )
        bounds = [(0.0, 5.0), (0.1, 5.0)]

        rho = float(self.parametros["rho"]["entry"].get())
        c = float(self.parametros["c"]["entry"].get())
        L = float(self.parametros["L"]["entry"].get())
        D = float(self.parametros["D"]["entry"].get())

        def objetivo(vars_):
            if self.scenario.get() == "rigido":
                t, P, Q = modelo_inercial_rigido(
                    rho, L, D, vars_[0], vars_[1], 
                    t_total=5.0, dt=0.01, P_inicial=self.P_inicial*1000
                )
                max_p = np.max(P)
            else:
                v_ini = vars_[0]
                t_cie = vars_[1]
                t, hup, hdown, _ = metodo_caracteristicas_avanzado(
                    rho, c, L, D,
                    v_inicial=v_ini,
                    t_cierre_valvula=t_cie,
                    scenario=self.scenario.get(),
                    H_up=self.H_up,
                    H_down=self.H_down,
                    t_total=5.0,
                    dt=0.01
                )
                g = 9.81
                p_up = rho * g * hup
                p_down = rho * g * hdown
                max_p = max(np.max(p_up), np.max(p_down))
                return max_p
        
        result = differential_evolution(objetivo, bounds, maxiter=20, popsize=15, tol=1e-3)
        if result.success:
            v_opt_inicial, t_opt_cierre = result.x
            pres_min = result.fun
            msg = (f"Optimización finalizada con éxito.\n\n"
                   f"Vel. inicial óptima: {v_opt_inicial:.3f} m/s\n"
                   f"Tiempo de cierre óptimo: {t_opt_cierre:.3f} s\n"
                   f"Presión máxima final: {pres_min:.2f} Pa\n")
            messagebox.showinfo("Resultado de la Optimización", msg)
        else:
            messagebox.showwarning(
                "Optimización incompleta",
                "El optimizador no convergió a una solución satisfactoria."
            )

    def dibujar_red_epanet(self):
        if not self.water_hammer_analyzer:
            return
            
        self.canvas.delete("all")
        network = self.water_hammer_analyzer.network
        
        coords = {}
        for node in network.nodes:
            coords[node.uid] = node.coordinates
        
        coords = {k: v for k, v in coords.items() if v is not None}
        
        if not coords:
            messagebox.showerror("Error", "El archivo .inp no contiene coordenadas de nodos")
            return
        
        x_vals = [v[0] for v in coords.values()]
        y_vals = [v[1] for v in coords.values()]
        min_x, max_x = min(x_vals), max(x_vals)
        min_y, max_y = min(y_vals), max(y_vals)
        
        canvas_width = 800
        canvas_height = 600
        scale_x = canvas_width / (max_x - min_x) if (max_x - min_x) != 0 else 1
        scale_y = canvas_height / (max_y - min_y) if (max_y - min_y) != 0 else 1
        
        for pipe in network.pipes:
            start_node = pipe.from_node.uid
            end_node = pipe.to_node.uid
            if start_node in coords and end_node in coords:
                x1 = (coords[start_node][0] - min_x) * scale_x
                y1 = (coords[start_node][1] - min_y) * scale_y
                x2 = (coords[end_node][0] - min_x) * scale_x
                y2 = (coords[end_node][1] - min_y) * scale_y
                self.canvas.create_line(x1, y1, x2, y2, fill="blue", width=2)
        
        for node_id, (x, y) in coords.items():
            x_scaled = (x - min_x) * scale_x
            y_scaled = (y - min_y) * scale_y
            self.canvas.create_oval(
                x_scaled-5, y_scaled-5, x_scaled+5, y_scaled+5,
                fill="red", outline="black"
            )
            self.canvas.create_text(
                x_scaled, y_scaled-10,
                text=node_id, fill="green", font=("Arial", 8)
            )
        
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

if __name__ == "__main__":
    app = AppGolpeAriete()
    app.mainloop()
